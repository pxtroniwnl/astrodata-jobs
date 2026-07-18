"""Extract text from PDF and DOCX files."""
from __future__ import annotations

import io


def extract_text(file_bytes: bytes, filename: str) -> str:
    """Extract plain text from a PDF, DOCX or TXT file.

    El tipo real se detecta por el contenido (magic bytes), no por la
    extensión: es común subir un DOCX renombrado a .pdf (o viceversa).
    La extensión solo se usa como respaldo cuando el contenido no es
    reconocible. Cualquier archivo ilegible produce ValueError con un
    mensaje claro para el usuario.
    """
    head = file_bytes[:1024]
    if head.lstrip().startswith(b"%PDF"):
        return _extract_pdf(file_bytes)
    if head.startswith(b"PK\x03\x04"):
        return _extract_docx(file_bytes)

    lower = filename.lower()
    if lower.endswith(".pdf"):
        return _extract_pdf(file_bytes)
    if lower.endswith(".docx"):
        return _extract_docx(file_bytes)
    if lower.endswith(".txt"):
        return file_bytes.decode("utf-8", errors="replace")
    raise ValueError(f"Formato no soportado: {filename}. Usa PDF, DOCX o TXT.")


def _extract_pdf(data: bytes) -> str:
    from pypdf import PdfReader

    try:
        reader = PdfReader(io.BytesIO(data))
        pages = []
        for page in reader.pages:
            text = page.extract_text()
            if text:
                pages.append(text)
    except Exception as e:
        raise ValueError(
            "No se pudo leer el PDF. Verifica que el archivo no esté dañado "
            "ni protegido con contraseña, o exporta tu CV de nuevo a PDF."
        ) from e
    return "\n\n".join(pages)


def _extract_docx(data: bytes) -> str:
    from docx import Document

    try:
        doc = Document(io.BytesIO(data))
        parts = []
        for para in doc.paragraphs:
            if para.text.strip():
                parts.append(para.text)
        for table in doc.tables:
            for row in table.rows:
                cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
                if cells:
                    parts.append(" | ".join(cells))
    except Exception as e:
        raise ValueError(
            "No se pudo leer el DOCX. Verifica que el archivo no esté dañado, "
            "o guárdalo de nuevo desde tu editor de texto."
        ) from e
    return "\n".join(parts)
